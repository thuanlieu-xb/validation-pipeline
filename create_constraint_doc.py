from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)

# Title
title = doc.add_heading("TAI LIEU HUAN LUYEN BO SUNG", 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub = doc.add_heading("Ma tran rang buoc: Business Unit x Billing Journal x Thue dau vao x Deliver To", 2)
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph(
    "Tai lieu nay dinh nghia chinh xac Billing Journal, Thue dau vao (Tax Input) va Deliver To "
    "tuong ung voi tung Business Unit trong he thong Odoo ERP. "
    "Dung ket hop voi PO_Validation_Training_Document.docx."
)

# 1. Quy tac tong quat
doc.add_heading("1. Quy tac tong quat", 1)
rules = [
    "Billing Journal phai khop CHINH XAC voi Business Unit cua PO.",
    "Thue dau vao kiem tra bang ILIKE '%KEYWORD%' - KHONG kiem tra % thue cu the.",
    "Deliver To phai dung loai kho tuong ung voi Journal da chon.",
    "Journal ALL BU (CCDC/VP/Tong hop/Hang khac) co the dung cho moi BU.",
    "XAY LAP DAC BIET: Billing Journal BAT BUOC la 'Hoa don mua vat lieu cho thi cong DA'.",
    "'Nhat ky vay' (TC01) KHONG duoc xuat hien tren PO mua hang thong thuong.",
    "GTGT TSCD (CCDC/TSCD): ap dung cho TAT CA BU — bat ky linh vuc nao cung co the mua tai san co dinh.",
    "GTGT Khac: ap dung cho TAT CA BU — hang hoa, dich vu khong thuoc nhom khac.",
    "GTGT Nhap khau: ap dung cho TAT CA BU — bat ky BU nao cung co the nhap khau hang hoa.",
    "GTGT HHDV: ap dung cho TAT CA BU — mua hang hoa dich vu thong thuong.",
]
for r in rules:
    doc.add_paragraph(r, style="List Bullet")

# 2. Ma tran rang buoc
doc.add_heading("2. Ma tran rang buoc chi tiet", 1)

headers = ["Hoat dong", "BU Code", "Business Unit", "Billing Journal (CHINH XAC)", "Deliver To", "Thue dau vao (keyword)"]
rows_data = [
    ["HOAT DONG DAU TU",     "DT01", "Dau tu: Xay lap",                    "Hoa don mua hang cho hoat dong dau tu",                      "Vat tu / NVL",         "GTGT Dau tu"],
    ["",                     "DT02", "Dau tu: Bat Dong San",               "Hoa don mua hang - KDDV",                                    "N/A",                  "N/A"],
    ["",                     "DT03", "Dau tu: Cho thue Van phong, BDS",    "Hoa don mua hang - KDDV",                                    "N/A",                  "N/A"],
    ["",                     "DT04", "Dau tu: Nang luong",                  "Chua co - can xac nhan",                                     "N/A",                  "N/A"],
    ["HOAT DONG DICH VU",    "DV01", "Dich vu: Tu Van Thiet Ke",           "Chua co - can xac nhan",                                     "N/A",                  "N/A"],
    ["",                     "DV02", "Dich vu: Tu Van Nguon nhan luc",     "Chua co - can xac nhan",                                     "N/A",                  "N/A"],
    ["HOAT DONG KINH DOANH", "KD01", "Kinh doanh: Vat tu xay dung",       "Hoa don mua hang hoa KDVT\nHoac: Hoa don mua hang - KDVT",   "Hang hoa",             "GTGT KDVT"],
    ["",                     "KD02", "Kinh doanh: Nong san",               "N/A",                                                        "N/A",                  "N/A"],
    ["",                     "KD03", "Kinh doanh: Nhien lieu",             "Chua co - can xac nhan",                                     "N/A",                  "N/A"],
    ["HOAT DONG KHAC",       "KK99", "Khac: Hoat dong khac",              "N/A",                                                        "N/A",                  "N/A"],
    ["HOAT DONG TAI CHINH",  "TC01", "Tai Chinh: Hoat dong tai chinh",    "Nhat ky vay (CHI dung TC01 - KHONG dung cho PO mua hang)",   "N/A",                  "N/A"],
    ["HOAT DONG XAY LAP",    "XL01", "Xay lap: Dan dung",                 "BAT BUOC: Hoa don mua vat lieu cho thi cong DA",             "Vat tu / NVL",         "GTGT Xay Lap"],
    ["",                     "XL02", "Xay lap: Cong nghiep",               "BAT BUOC: Hoa don mua vat lieu cho thi cong DA",             "Vat tu / NVL",         "GTGT Xay Lap"],
    ["",                     "XL03", "Xay lap: Giao thong & Ha tang",      "BAT BUOC: Hoa don mua vat lieu cho thi cong DA",             "Vat tu / NVL",         "GTGT Xay Lap"],
    ["HOAT DONG CHUNG",      "ALL",  "Ap dung toan bo BU",                "Hoa don mua CCDC",                                           "CCDC",                 "GTGT TSCD — ap dung TAT CA BU"],
    ["",                     "ALL",  "",                                    "Hoa don mua do dung van phong",                              "Van phong / Noi bo",   "GTGT HHDV — ap dung TAT CA BU"],
    ["",                     "ALL",  "",                                    "Hoa don mua hang tong hop",                                   "N/A",                  "GTGT Khac — ap dung TAT CA BU"],
    ["",                     "ALL",  "",                                    "Hoa don mua hang khac",                                       "N/A",                  "GTGT Nhap khau — ap dung TAT CA BU"],
]

table = doc.add_table(rows=1 + len(rows_data), cols=len(headers))
table.style = "Table Grid"

hdr_row = table.rows[0]
for i, h in enumerate(headers):
    cell = hdr_row.cells[i]
    cell.text = h
    cell.paragraphs[0].runs[0].bold = True
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_cell_bg(cell, "1F4E79")
    cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

for r_idx, row_data in enumerate(rows_data):
    row = table.rows[r_idx + 1]
    bu_code = row_data[1]
    for c_idx, val in enumerate(row_data):
        cell = row.cells[c_idx]
        cell.text = val
        if bu_code in ("XL01", "XL02", "XL03"):
            set_cell_bg(cell, "FFF2CC")
            if c_idx == 3 and "BAT BUOC" in val:
                for run in cell.paragraphs[0].runs:
                    run.bold = True
                    run.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)

# 3. Quy tac dac biet Xay Lap
doc.add_heading("3. Quy tac dac biet - XAY LAP (XL01, XL02, XL03)", 1)
p = doc.add_paragraph()
r = p.add_run("BILLING JOURNAL BAT BUOC CHO TAT CA BU XAY LAP:")
r.bold = True
r.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)

xl_rules = [
    'Journal PHAI LA: "Hoa don mua vat lieu cho thi cong DA"',
    'KHONG dung: "Hoa don NTP hoat dong tong hop" cho PO mua vat lieu thi cong',
    'KHONG dung: "Hoa don khau tru chu dau tu hoat dong xay lap" cho PO mua hang',
    'Deliver To: BAT BUOC la kho Vat tu / NVL (khong phai CCDC, khong phai Hang hoa)',
    'Tax keyword: ILIKE "%GTGT Xay Lap%" - Khong kiem tra %',
    'Analytic Distribution: phai thuoc BU XL* (XL01, XL02, XL03)',
    'Subcontractor Contract Project: thuong co du lieu -> +3 diem confidence',
    'Neu Journal la "Hoa don NTP hoat dong tong hop" tren PO mua vat lieu -> LOI [Critical]',
]
for rule in xl_rules:
    doc.add_paragraph(rule, style="List Bullet")

# 4. Mapping nhanh
doc.add_heading("4. Bang mapping nhanh Journal -> Tax keyword", 1)

doc.add_paragraph(
    "LUU Y QUAN TRONG: Cac loai thue GTGT TSCD, GTGT Khac, GTGT Nhap khau, GTGT HHDV "
    "co the xuat hien tren BAT KY BU nao. Day la cac journal 'ALL BU' — chi rang buoc "
    "o Deliver To va Tax tuong ung, KHONG rang buoc voi BU cu the."
)

mapping = [
    ("Hoa don mua vat lieu cho thi cong DA",   "GTGT Xay Lap",   "Vat tu / NVL",       "XL01, XL02, XL03 (BAT BUOC)"),
    ("Hoa don mua hang cho hoat dong dau tu",  "GTGT Dau tu",    "Vat tu / NVL",       "DT01"),
    ("Hoa don mua hang hoa KDVT",              "GTGT KDVT",      "Hang hoa",           "KD01"),
    ("Hoa don mua hang - KDVT",                "GTGT KDVT",      "Hang hoa",           "KD01"),
    ("Hoa don mua CCDC",                       "GTGT TSCD",      "CCDC",               "TAT CA BU (ALL BU)"),
    ("Hoa don mua do dung van phong",          "GTGT HHDV",      "Van phong / Noi bo", "TAT CA BU (ALL BU)"),
    ("Hoa don mua hang tong hop",              "GTGT Khac",      "N/A",                "TAT CA BU (ALL BU)"),
    ("Hoa don mua hang khac",                  "GTGT Nhap khau", "N/A",                "TAT CA BU (ALL BU)"),
    ("Nhat ky vay [CANH BAO]",                 "N/A",            "N/A",                "TC01 ONLY - KHONG dung cho PO mua hang"),
]

t2 = doc.add_table(rows=1 + len(mapping), cols=4)
t2.style = "Table Grid"
for i, h in enumerate(["Billing Journal", "Tax keyword (ILIKE)", "Deliver To", "BU ap dung"]):
    cell = t2.rows[0].cells[i]
    cell.text = h
    cell.paragraphs[0].runs[0].bold = True
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_cell_bg(cell, "1F4E79")
    cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

for r_idx, (journal, tax, deliver, bu) in enumerate(mapping):
    row = t2.rows[r_idx + 1]
    is_xl = "XL01" in bu
    for c_idx, val in enumerate([journal, tax, deliver, bu]):
        cell = row.cells[c_idx]
        cell.text = val
        if is_xl:
            set_cell_bg(cell, "FFF2CC")

# 5. Cac loai thue ALL BU
doc.add_heading("5. Cac loai thue ap dung cho TAT CA Business Unit", 1)
doc.add_paragraph(
    "Cac loai thue duoi day KHONG bi rang buoc voi BU cu the. "
    "Bat ky hoat dong nao (XL, DT, KD, DV...) cung co the su dung, "
    "mien la Billing Journal va Deliver To phu hop:"
)

allbu_taxes = [
    ("GTGT TSCD (ILIKE '%GTGT TSCD%')",
     "Mua tai san co dinh (xe o to, may moc, thiet bi...) — gia tri >= 30 trieu, su dung >= 1 nam. "
     "Ap dung cho tat ca BU: XL, DT, KD, DV, KK... "
     "Journal tuong ung: 'Hoa don mua CCDC'. Deliver To: kho CCDC."),
    ("GTGT Khac (ILIKE '%GTGT Khac%')",
     "Hang hoa dich vu khong thuoc nhom chuyen biet. "
     "Ap dung cho tat ca BU. Journal: 'Hoa don mua hang tong hop'."),
    ("GTGT Nhap khau (ILIKE '%GTGT Nhap khau%')",
     "Hang hoa nhap khau tu nuoc ngoai. "
     "Ap dung cho tat ca BU. Journal: 'Hoa don mua hang khac'."),
    ("GTGT HHDV (ILIKE '%GTGT HHDV%')",
     "Hang hoa dich vu thong thuong mua trong nuoc. "
     "Ap dung cho tat ca BU. Journal: 'Hoa don mua do dung van phong'."),
]

for tax_name, desc in allbu_taxes:
    p = doc.add_paragraph(style="List Bullet")
    run = p.add_run(tax_name + ": ")
    run.bold = True
    p.add_run(desc)

doc.add_paragraph(
    "VI DU: PO cua XL01 mua xe o to (TSCD) -> dung 'Hoa don mua CCDC' + GTGT TSCD -> PASS. "
    "PO cua KD01 mua thiet bi van phong nhap khau -> dung 'Hoa don mua hang khac' + GTGT Nhap khau -> PASS. "
    "Day KHONG phai loi cross-BU."
)

path = "data/BU_Journal_Tax_Constraints.docx"
doc.save(path)
print(f"Da tao: {path}")
